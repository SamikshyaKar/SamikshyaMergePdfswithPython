import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
from pathlib import Path
import sys

try:
    from PyPDF2 import PdfMerger, PdfReader
    import docx
    from PIL import Image
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.utils import ImageReader
except ImportError as e:
    print(f"Missing required library: {e}")
    print("Please install requirements: pip install -r requirements.txt")
    sys.exit(1)


class FileMergerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Universal File Merger")
        self.root.geometry("700x600")
        self.root.resizable(True, True)
        
        self.selected_files = []
        self.output_path = ""
        
        self.setup_ui()
        
    def setup_ui(self):
        # Title
        title_label = tk.Label(
            self.root, 
            text="Universal File Merger", 
            font=("Arial", 18, "bold"),
            pady=10
        )
        title_label.pack()
        
        # Description
        desc_label = tk.Label(
            self.root,
            text="Merge PDF, Word, and Image files into a single PDF",
            font=("Arial", 10),
            fg="gray"
        )
        desc_label.pack(pady=5)
        
        # File selection frame
        file_frame = tk.Frame(self.root, padx=20, pady=10)
        file_frame.pack(fill=tk.BOTH, expand=True)
        
        # Add files button
        add_btn = tk.Button(
            file_frame,
            text="Add Files",
            command=self.add_files,
            bg="#4CAF50",
            fg="white",
            font=("Arial", 11, "bold"),
            padx=20,
            pady=5
        )
        add_btn.pack(pady=10)
        
        # File list with scrollbar
        list_frame = tk.Frame(file_frame)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.file_listbox = tk.Listbox(
            list_frame,
            yscrollcommand=scrollbar.set,
            font=("Arial", 10),
            selectmode=tk.EXTENDED
        )
        self.file_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=self.file_listbox.yview)
        
        # Remove selected button
        remove_btn = tk.Button(
            file_frame,
            text="Remove Selected",
            command=self.remove_selected,
            bg="#f44336",
            fg="white",
            font=("Arial", 10),
            padx=15,
            pady=5
        )
        remove_btn.pack(pady=5)
        
        # Output path frame
        output_frame = tk.Frame(self.root, padx=20, pady=10)
        output_frame.pack(fill=tk.X)
        
        tk.Label(
            output_frame,
            text="Output File:",
            font=("Arial", 10, "bold")
        ).pack(anchor=tk.W)
        
        output_path_frame = tk.Frame(output_frame)
        output_path_frame.pack(fill=tk.X, pady=5)
        
        self.output_entry = tk.Entry(
            output_path_frame,
            font=("Arial", 10),
            state="readonly"
        )
        self.output_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))
        
        browse_btn = tk.Button(
            output_path_frame,
            text="Browse",
            command=self.browse_output,
            font=("Arial", 10),
            padx=10
        )
        browse_btn.pack(side=tk.RIGHT)
        
        # Progress bar
        self.progress = ttk.Progressbar(
            self.root,
            mode='indeterminate',
            length=400
        )
        self.progress.pack(pady=10)
        
        # Merge button
        merge_btn = tk.Button(
            self.root,
            text="Merge Files",
            command=self.merge_files,
            bg="#2196F3",
            fg="white",
            font=("Arial", 12, "bold"),
            padx=30,
            pady=10
        )
        merge_btn.pack(pady=20)
        
    def add_files(self):
        files = filedialog.askopenfilenames(
            title="Select files to merge",
            filetypes=[
                ("All Supported", "*.pdf;*.docx;*.doc;*.jpg;*.jpeg;*.png;*.gif;*.bmp;*.tiff"),
                ("PDF files", "*.pdf"),
                ("Word documents", "*.docx;*.doc"),
                ("Image files", "*.jpg;*.jpeg;*.png;*.gif;*.bmp;*.tiff"),
                ("All files", "*.*")
            ]
        )
        
        for file in files:
            if file not in self.selected_files:
                self.selected_files.append(file)
                self.file_listbox.insert(tk.END, os.path.basename(file))
        
    def remove_selected(self):
        selected_indices = self.file_listbox.curselection()
        for index in reversed(selected_indices):
            self.file_listbox.delete(index)
            del self.selected_files[index]
    
    def browse_output(self):
        output_file = filedialog.asksaveasfilename(
            title="Save merged file as",
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        
        if output_file:
            self.output_path = output_file
            self.output_entry.config(state="normal")
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, output_file)
            self.output_entry.config(state="readonly")
    
    def convert_image_to_pdf(self, image_path, output_pdf_path):
        """Convert an image file to PDF"""
        try:
            img = Image.open(image_path)
            # Convert RGBA to RGB if necessary
            if img.mode == 'RGBA':
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                rgb_img.paste(img, mask=img.split()[3])
                img = rgb_img
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Calculate dimensions to fit on page
            img_width, img_height = img.size
            page_width, page_height = A4
            
            # Scale to fit page while maintaining aspect ratio
            scale_w = page_width / img_width
            scale_h = page_height / img_height
            scale = min(scale_w, scale_h)
            
            new_width = img_width * scale
            new_height = img_height * scale
            
            # Center image on page
            x_offset = (page_width - new_width) / 2
            y_offset = (page_height - new_height) / 2
            
            # Create PDF
            c = canvas.Canvas(output_pdf_path, pagesize=A4)
            c.drawImage(ImageReader(img), x_offset, y_offset, width=new_width, height=new_height)
            c.save()
            return True
        except Exception as e:
            print(f"Error converting image {image_path}: {e}")
            return False
    
    def convert_word_to_pdf(self, word_path, output_pdf_path):
        """Convert Word document to PDF using python-docx and reportlab"""
        try:
            doc = docx.Document(word_path)
            
            c = canvas.Canvas(output_pdf_path, pagesize=letter)
            width, height = letter
            y_position = height - 50
            line_height = 14
            margin = 50
            
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if not text.strip():
                    y_position -= line_height
                    continue
                
                # Simple text wrapping
                words = text.split()
                line = ""
                for word in words:
                    test_line = line + word + " " if line else word + " "
                    text_width = c.stringWidth(test_line, "Helvetica", 10)
                    
                    if text_width > (width - 2 * margin) and line:
                        c.drawString(margin, y_position, line)
                        y_position -= line_height
                        line = word + " "
                        
                        if y_position < margin:
                            c.showPage()
                            y_position = height - 50
                    else:
                        line = test_line
                
                if line:
                    c.drawString(margin, y_position, line)
                    y_position -= line_height
                    
                    if y_position < margin:
                        c.showPage()
                        y_position = height - 50
            
            # Handle tables
            for table in doc.tables:
                y_position -= line_height * 2
                if y_position < margin:
                    c.showPage()
                    y_position = height - 50
                
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if len(row_text) > 100:
                        row_text = row_text[:100] + "..."
                    c.drawString(margin, y_position, row_text)
                    y_position -= line_height
                    
                    if y_position < margin:
                        c.showPage()
                        y_position = height - 50
            
            c.save()
            return True
        except Exception as e:
            print(f"Error converting Word document {word_path}: {e}")
            return False
    
    def merge_files(self):
        if not self.selected_files:
            messagebox.showwarning("Warning", "Please select at least one file to merge.")
            return
        
        if not self.output_path:
            messagebox.showwarning("Warning", "Please select an output file path.")
            return
        
        try:
            self.progress.start()
            self.root.update()
            
            # Create temporary directory for converted files
            temp_dir = Path(self.output_path).parent / "temp_merge"
            temp_dir.mkdir(exist_ok=True)
            
            pdf_files = []
            
            # Process each file
            for i, file_path in enumerate(self.selected_files):
                file_ext = Path(file_path).suffix.lower()
                temp_pdf = temp_dir / f"temp_{i}.pdf"
                
                if file_ext == '.pdf':
                    # PDF files can be merged directly
                    pdf_files.append(file_path)
                elif file_ext in ['.docx', '.doc']:
                    # Convert Word to PDF
                    if self.convert_word_to_pdf(file_path, str(temp_pdf)):
                        pdf_files.append(str(temp_pdf))
                    else:
                        messagebox.showerror("Error", f"Failed to convert {os.path.basename(file_path)}")
                        return
                elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                    # Convert image to PDF
                    if self.convert_image_to_pdf(file_path, str(temp_pdf)):
                        pdf_files.append(str(temp_pdf))
                    else:
                        messagebox.showerror("Error", f"Failed to convert {os.path.basename(file_path)}")
                        return
                else:
                    messagebox.showwarning("Warning", f"Unsupported file type: {file_path}")
            
            # Merge all PDFs
            if pdf_files:
                merger = PdfMerger()
                for pdf_file in pdf_files:
                    merger.append(pdf_file)
                
                merger.write(self.output_path)
                merger.close()
                
                # Clean up temporary files
                import shutil
                shutil.rmtree(temp_dir, ignore_errors=True)
                
                self.progress.stop()
                messagebox.showinfo("Success", f"Files merged successfully!\nSaved to: {self.output_path}")
            else:
                self.progress.stop()
                messagebox.showerror("Error", "No valid files to merge.")
                
        except Exception as e:
            self.progress.stop()
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
            import traceback
            traceback.print_exc()


def main():
    root = tk.Tk()
    app = FileMergerApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()

